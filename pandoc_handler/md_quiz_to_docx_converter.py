# -*- coding: utf-8 -*-
# md_quiz_to_docx_converter.py

"""
Script para reestructurar archivos Markdown de cuestionarios.

Dependencias:
- panflute: Una biblioteca para crear filtros de Pandoc en Python.
  Para instalarla, ejecuta: pip install panflute

Descripción:
Este script procesa archivos Markdown (.md) ubicados en una carpeta 'Temporales'.
La lógica de transformación es la siguiente:
1. Lee un archivo Markdown que contiene bloques de texto seguidos de preguntas
   de opción múltiple basadas en ese texto.
2. Identifica estos bloques de texto (contexto) y los bloques de preguntas
   (listas ordenadas).
3. Para cada pregunta en una lista, mueve el bloque de texto de contexto
   correspondiente al interior de la pregunta.
4. Antecede al texto movido la frase "De acuerdo al texto responde:" seguida
   de dos saltos de línea.
5. Conserva todo el formato de texto enriquecido (negritas, cursivas,
   resaltados, imágenes, etc.) gracias al uso de un AST (Abstract Syntax Tree)
   manejado por panflute/pandoc.
"""

import os
import subprocess
import panflute as pf
from typing import List, Optional

from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_COLOR_INDEX

import re
from pathlib import Path

class DocxPostProcessor:
    """
    Clase encargada de la limpieza y ajuste fino de archivos DOCX.
    Usa inspección profunda de definiciones de numeración (Deep Inspection).
    """

    def _is_bullet_list(self, doc, numId_val):
        """
        Investiga en el numbering.xml si un ID específico corresponde a una viñeta.
        Retorna True si es viñeta (bullet), False si es número (decimal, etc).
        """
        try:
            # Accedemos a la parte de numeración del documento
            numbering_part = doc.part.numbering_part
            if numbering_part is None:
                return False

            # 1. Buscar el elemento <w:num> que coincida con el numId del párrafo
            # Esto nos dará el abstractNumId
            num_element = numbering_part._element.find(f'.//w:num[@w:numId="{numId_val}"]', numbering_part._element.nsmap)
            if num_element is None:
                return False

            abstract_num_id_element = num_element.find(qn('w:abstractNumId'))
            if abstract_num_id_element is None:
                return False
            
            abstract_num_val = abstract_num_id_element.get(qn('w:val'))

            # 2. Buscar la definición abstracta <w:abstractNum> usando el ID encontrado
            abstract_num = numbering_part._element.find(f'.//w:abstractNum[@w:abstractNumId="{abstract_num_val}"]', numbering_part._element.nsmap)
            if abstract_num is None:
                return False

            # 3. Buscar el nivel 0 (<w:lvl w:ilvl="0">) para ver cómo está formateado
            lvl = abstract_num.find(f'.//w:lvl[@w:ilvl="0"]', numbering_part._element.nsmap)
            if lvl is None:
                return False

            # 4. Verificar el formato de número (<w:numFmt w:val="...">)
            num_fmt = lvl.find(qn('w:numFmt'))
            if num_fmt is None:
                return False

            fmt_val = num_fmt.get(qn('w:val'))
            
            # Si el formato es 'bullet', es lo que queremos eliminar.
            # Si es 'decimal', 'lowerLetter', etc., lo conservamos.
            return fmt_val == 'bullet'

        except Exception:
            # En caso de error de lectura XML, asumimos que NO es bullet para no romper nada importante
            return False

    def remove_bullets_keep_indent(self, file_path: str):
        try:
            doc = Document(file_path)
            modified = False
            
            for paragraph in doc.paragraphs:
                # 1. Obtener acceso al XML del párrafo
                p_pr = paragraph._element.find(qn('w:pPr'))
                
                if p_pr is not None:
                    # 2. Ver si tiene propiedades de numeración
                    num_pr = p_pr.find(qn('w:numPr'))
                    
                    if num_pr is not None:
                        # 3. Extraer el ID de la lista
                        num_id_element = num_pr.find(qn('w:numId'))
                        if num_id_element is not None:
                            val = num_id_element.get(qn('w:val'))
                            
                            # 4. INSPECCIÓN PROFUNDA:
                            # Solo procedemos si confirmamos que este ID es una VIÑETA.
                            if self._is_bullet_list(doc, val):
                                
                                # Capturar indentación actual antes de borrar
                                current_indent = paragraph.paragraph_format.left_indent
                                if current_indent is None:
                                    current_indent = Pt(36) # Default ~1.27cm
                                
                                # ELIMINAR LA VIÑETA (Borrar nodo numPr)
                                p_pr.remove(num_pr)
                                
                                # RESTAURAR INDENTACIÓN
                                paragraph.paragraph_format.left_indent = current_indent
                                
                                modified = True

            if modified:
                doc.save(file_path)
                print(f"    -> Post-procesamiento: Viñetas eliminadas (Numeración conservada) en {file_path}")
            else:
                print(f"    -> Post-procesamiento: No se requirieron cambios en {file_path}")

        except Exception as e:
            print(f"    -> ERROR en post-procesamiento: {e}")

    def _force_list_level_format(self, doc, num_id, level=1, format_type='upperLetter'):
        """
        Inspecciona el XML para encontrar la definición abstracta de la lista y 
        forzar el formato de número (ej. A, B, C) y EL INICIO EN 1 para un nivel específico.
        """
        try:
            numbering_part = doc.part.numbering_part
            if numbering_part is None:
                return

            # 1. Encontrar el elemento num (<w:num>) usando el numId actual
            num_element = numbering_part._element.find(f'.//w:num[@w:numId="{num_id}"]', numbering_part._element.nsmap)
            if num_element is None:
                return
            
            # Obtener el abstractNumId asociado
            abstract_num_id_element = num_element.find(qn('w:abstractNumId'))
            if abstract_num_id_element is None:
                return
            abstract_num_val = abstract_num_id_element.get(qn('w:val'))

            # 2. Encontrar la definición abstracta (<w:abstractNum>)
            abstract_num = numbering_part._element.find(f'.//w:abstractNum[@w:abstractNumId="{abstract_num_val}"]', numbering_part._element.nsmap)
            if abstract_num is None:
                return

            # 3. Modificar la Definición Abstracta (Afecta a todas las listas de este tipo)
            lvl = abstract_num.find(f'.//w:lvl[@w:ilvl="{level}"]', numbering_part._element.nsmap)
            if lvl is not None:
                # A. Forzar formato de letra (A., B., C.)
                num_fmt = lvl.find(qn('w:numFmt'))
                if num_fmt is not None:
                    num_fmt.set(qn('w:val'), format_type)
                
                # B. Forzar inicio en 1 (<w:start w:val="1"/>) - CORRECCIÓN CLAVE
                start_node = lvl.find(qn('w:start'))
                if start_node is not None:
                    start_node.set(qn('w:val'), '1')

            # 4. Modificar la Instancia Específica (Override) - CORRECCIÓN CLAVE
            # A veces la lista específica tiene una anulación (override) que fuerza el inicio en 4.
            # Debemos buscar <w:lvlOverride w:ilvl="1"> y cambiar <w:startOverride w:val="1">
            lvl_override = num_element.find(f'.//w:lvlOverride[@w:ilvl="{level}"]', numbering_part._element.nsmap)
            if lvl_override is not None:
                start_override = lvl_override.find(qn('w:startOverride'))
                if start_override is not None:
                    start_override.set(qn('w:val'), '1')

        except Exception as e:
            print(f"    -> Warning: No se pudo forzar el formato de lista: {e}")

    def convert_text_options_to_list(self, file_path: str):
        doc = Document(file_path)
        modified = False
        prefix_pattern = re.compile(r'^\s*([A-Da-d])(\.|\))\s+') 
        last_valid_num_id = None

        for paragraph in doc.paragraphs:
            # 1. Capturar contexto de lista
            p_pr = paragraph._element.find(qn('w:pPr'))
            if p_pr is not None:
                num_pr = p_pr.find(qn('w:numPr'))
                if num_pr is not None:
                    num_id_node = num_pr.find(qn('w:numId'))
                    if num_id_node is not None:
                        val_str = num_id_node.get(qn('w:val'))
                        ilvl_node = num_pr.find(qn('w:ilvl'))
                        if (ilvl_node is None or ilvl_node.get(qn('w:val')) == '0') and val_str:
                            last_valid_num_id = val_str

            # 2. Identificar si es una opción (A, B, C...)
            text = paragraph.text
            match = prefix_pattern.match(text)
            
            if match and last_valid_num_id:
                # Detectar si debe ser resaltada (basado en si algún run es negrita originalmente)
                is_correct_answer = any(run.bold for run in paragraph.runs)
                
                # --- LIMPIEZA TOTAL DE FORMATO HEREDADO ---
                if p_pr is not None:
                    p_style = p_pr.find(qn('w:pStyle'))
                    if p_style is not None:
                        p_pr.remove(p_style)

                # --- TRATAMIENTO QUIRÚRGICO DE RUNS ---
                chars_to_remove = len(match.group(0))
                for run in paragraph.runs:
                    # Quitamos la negrita de forma radical borrando la propiedad
                    run.bold = None 
                    
                    # Si es la respuesta correcta, aplicamos resaltado al texto
                    if is_correct_answer:
                        run.font.highlight_color = WD_COLOR_INDEX.YELLOW
                        # Acceso directo al XML para asegurar que no hay rPr de negrita
                        rPr = run._element.get_or_add_rPr()
                        for tag in [qn('w:b'), qn('w:bCs')]:
                            node = rPr.find(tag)
                            if node is not None:
                                rPr.remove(node)

                    # Eliminación del prefijo
                    if chars_to_remove > 0:
                        run_text = run.text
                        if len(run_text) <= chars_to_remove:
                            chars_to_remove -= len(run_text)
                            run.text = ""
                        else:
                            run.text = run_text[chars_to_remove:]
                            chars_to_remove = 0

                # --- RE-VINCULACIÓN A LA LISTA E INDENTACIÓN (PRIMERO) ---
                p_pr = paragraph._element.get_or_add_pPr()
                num_pr = p_pr.get_or_add_numPr()
                num_id_elem = num_pr.get_or_add_numId()
                num_id_elem.val = int(last_valid_num_id)
                ilvl_elem = num_pr.get_or_add_ilvl()
                ilvl_elem.val = 1 
                
                self._force_list_level_format(doc, last_valid_num_id, level=1, format_type='upperLetter')
                
                # Aplicamos la indentación ANTES de agregar el rPr manual
                paragraph.paragraph_format.left_indent = Pt(54)
                paragraph.paragraph_format.first_line_indent = Pt(-18)

                # --- RESALTADO DE LA MARCA DE PÁRRAFO / VIÑETA (AL FINAL) ---
                if is_correct_answer:
                    # Como ya se aplicó la indentación, el append colocará el rPr al final (Orden XML correcto)
                    r_pr = p_pr.find(qn('w:rPr'))
                    
                    if r_pr is None:
                        r_pr = OxmlElement('w:rPr')
                        p_pr.append(r_pr)
                    
                    highlight = r_pr.find(qn('w:highlight'))
                    if highlight is None:
                        highlight = OxmlElement('w:highlight')
                        r_pr.append(highlight)
                    
                    highlight.set(qn('w:val'), 'yellow')

                modified = True

        if modified:
            doc.save(file_path)


class MdQuizToDocxConverter:
    """
    Clase que encapsula la lógica para reestructurar archivos Markdown
    de cuestionarios.
    """

    def __init__(self,
        source_folder: str = "Temporales",
        destination_folder: str = "TemporalesTextoAVoz",
        reuse_stimulus_input: bool = False):
        """
        Inicializa el restructurador.

        Args:
            source_folder (str): El nombre de la carpeta que contiene los
                                 archivos .md a procesar. Se espera que
                                 esté al mismo nivel que el script.
        """
        self.source_path = Path(source_folder)
        self.destination_path = Path(destination_folder)
        self.reuse_stimulus = reuse_stimulus_input

    def reorder_doc(self, doc: pf.Doc) -> pf.Doc:
        '''
        Reorganiza las secciones no numeradas en lista numerada
        '''
        
        new_content: List[pf.Block]=[]
        current_context_blocks: List[pf.Block] = []
        prefix_text = "De acuerdo al texto responde:"
        each_list_item_content: List[pf.ListItem]=[]
                
        for elem in doc._content:
            if isinstance(elem, pf.OrderedList):
                prefix_para = pf.Para(pf.Str(prefix_text))

                cloned_context = [b for b in current_context_blocks]
                each_item_content: List[pf.Block]=[]
                
                #print(cloned_context)

                for each_list_item in elem._content:
                    for each_item in each_list_item._content:
                        each_item_content.append(each_item)
                    each_list_item_content.append(pf.ListItem(prefix_para,*cloned_context,*each_item_content))
                    each_item_content=[]

                #new_content.append(pf.Para())
                current_context_blocks = []
            else:
                current_context_blocks.append(elem)
        
        ordered_list_content=pf.OrderedList(*each_list_item_content)
        new_content.append(ordered_list_content)
        #doc=pf.Doc(*new_content)
        doc.content=new_content
        return doc


    def _process_file(self, file_path: str, output_file_path: str):
        """
        Procesa un único archivo Markdown, aplicando la transformación.
        """
        print(f"Procesando archivo: {os.path.basename(file_path)}...")
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                original_content = f.read()
            #original_content.encode('utf-8-sig').strip()

            # 1. Convertir el texto de entrada a un objeto Doc (AST)
            doc = pf.convert_text(
                original_content,
                input_format='markdown+mark',
                output_format='panflute',
                standalone=True,
                extra_args=['--wrap=none']
            )
            
            # 2. Ordenar el Doc en memoria
            doc=self.reorder_doc(doc)

            # 3. Convertir el objeto Doc modificado de vuelta a texto Markdown
            final_markdown_text = pf.convert_text(
                doc,
                input_format='panflute',
                output_format='markdown+mark',
                extra_args=['--wrap=none']
            )

            # 4. Escribir la cadena de texto resultante en el archivo
            with open(output_file_path, 'w', encoding='utf-8') as f:
                f.write(final_markdown_text)
            print(f" -> Archivo transformado exitosamente.")

        except Exception as e:
            print(f" -> ERROR procesando el archivo {os.path.basename(file_path)}: {e}")

    def run(self):
        """
        Ejecuta el proceso de reestructuración para todos los archivos .md
        en la carpeta de origen.
        """
        if not os.path.isdir(self.source_path):
            print(f"El directorio de origen '{self.source_path}' no fue encontrado.")
            print("El script no procesará ningún archivo.")
            return

        print(f"Iniciando proceso en la carpeta: '{self.source_path}'")
        markdown_files_found = False
        for filename in os.listdir(self.source_path):
            if filename.lower().endswith(".md"):
                markdown_files_found = True
                full_path = os.path.join(self.source_path, filename)
                path_to_normal=os.path.join(self.destination_path,filename)
                #path_to_images=os.path.join(self.source_path,"Imagenes")
                full_path_output = os.path.join(self.destination_path, filename.replace(".md","-modificado.md"))
                
                command = [
                    "pandoc",
                    str(full_path),
                    "-o", str(path_to_normal.replace(".md",".docx")),
                    "--wrap=none",
                    f"--resource-path=.{os.pathsep}{self.source_path}"]
                
                subprocess.run(
                    command,
                    check=True,        # Lanza una excepción si el comando devuelve un código de error.
                    capture_output=False, # Captura la salida estándar y el error estándar.
                    text=True,
                    encoding='utf-8'
                    )
                
                # Instanciamos el procesador (puedes hacerlo fuera del loop si prefieres optimizar)
                docx_cleaner = DocxPostProcessor()

                # Aplicamos la limpieza al archivo docx recién creado
                docx_path = str(path_to_normal.replace(".md", ".docx"))
                docx_cleaner.remove_bullets_keep_indent(docx_path)
                #Convertir opciones A. B. C. a lista real
                docx_cleaner.convert_text_options_to_list(docx_path)

                if self.reuse_stimulus:

                    self._process_file(full_path,full_path_output)
                    #print(path_to_images)
                    #Convertir los archivos en docx

                    command = [
                        "pandoc",
                        str(full_path_output),
                        "-o",
                        str(full_path_output.replace(".md",".docx")),
                        "--wrap=none",
                        f"--resource-path=:{self.source_path}"]
                    #os.system("pandoc " + full_path_output + " -o " + full_path_output.replace(".md",".docx"))
                    subprocess.run(
                        command,
                        check=True,        # Lanza una excepción si el comando devuelve un código de error.
                        capture_output=True, # Captura la salida estándar y el error estándar.
                        text=True, # Decodifica la salida y el error como texto.
                        encoding='utf-8'
                        )
                    
                    # Aplicamos la limpieza al archivo docx recién creado
                    docx_path = str(full_path_output.replace(".md", ".docx"))
                    docx_cleaner.remove_bullets_keep_indent(docx_path)
                    #Convertir opciones A. B. C. a lista real
                    docx_cleaner.convert_text_options_to_list(docx_path)
        
        if not markdown_files_found:
            print("No se encontraron archivos .md en la carpeta.")
            
        print("Proceso completado.")


if __name__ == '__main__':
    FOLDER = "Temporales"
    
    try:
        
        restructurer = MdQuizToDocxConverter(source_folder=FOLDER,destination_folder=FOLDER)
        restructurer.run()
        
    except ImportError:
        print("\nERROR: La biblioteca 'panflute' no está instalada.")
        print("Por favor, ejecute: pip install panflute")
    except Exception as e:
        print(f"Ha ocurrido un error inesperado durante la ejecución: {e}")