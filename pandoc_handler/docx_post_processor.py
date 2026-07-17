from docx import Document
from docx.shared import Pt, Cm
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_COLOR_INDEX
from typing import Optional
import re

class DocxPostProcessor:
    """
    Clase encargada de la limpieza y ajuste fino de archivos DOCX.
    Usa inspección profunda de definiciones de numeración (Deep Inspection).
    """
    def __init__(self, target_font: Optional[str] = None, target_font_size: Optional[int] = None):
        self.target_font = target_font
        self.target_font_size = target_font_size

    def apply_global_font(self, file_path: str):
        # Si no hay fuente ni tamaño definido, no hacemos nada
        if not self.target_font and not self.target_font_size:
            return
            
        try:
            doc = Document(file_path)
            
            # 1. Modificar el estilo base 'Normal'
            style = doc.styles['Normal']
            if self.target_font:
                style.font.name = self.target_font
            if self.target_font_size:
                style.font.size = Pt(self.target_font_size)
            
            # 2. Modificar los valores por defecto a nivel de documento (Deep XML override)
            for rPr in doc.element.xpath('.//w:rPrDefault/w:rPr'):
                
                # A. Ajuste de familia tipográfica
                if self.target_font:
                    rFonts = rPr.find(qn('w:rFonts'))
                    if rFonts is None:
                        rFonts = OxmlElement('w:rFonts')
                        rPr.append(rFonts)
                    
                    rFonts.set(qn('w:ascii'), self.target_font)
                    rFonts.set(qn('w:hAnsi'), self.target_font)
                    rFonts.set(qn('w:cs'), self.target_font)
                
                # B. Ajuste de tamaño de fuente
                if self.target_font_size:
                    # En OOXML el tamaño se define en medios puntos (ej. 12pt = 24)
                    half_points = str(int(self.target_font_size * 2))
                    
                    # Para texto regular (sz)
                    sz = rPr.find(qn('w:sz'))
                    if sz is None:
                        sz = OxmlElement('w:sz')
                        rPr.append(sz)
                    sz.set(qn('w:val'), half_points)
                    
                    # Para texto complejo/símbolos (szCs)
                    szCs = rPr.find(qn('w:szCs'))
                    if szCs is None:
                        szCs = OxmlElement('w:szCs')
                        rPr.append(szCs)
                    szCs.set(qn('w:val'), half_points)
                
            doc.save(file_path)
            
            msg = f"    -> Post-procesamiento: Fuente ajustada"
            if self.target_font: msg += f" a '{self.target_font}'"
            if self.target_font_size: msg += f" tamaño {self.target_font_size}pt"
            print(f"{msg} en {file_path}")
            
        except Exception as e:
            print(f"    -> ERROR al ajustar formato global en {file_path}: {e}")

    def _is_bullet_list(self, doc, numId_val):
        """
        Investiga en el numbering.xml si un ID específico corresponde a una viñeta.
        Retorna True si es viñeta (bullet), False si es número (decimal, etc).
        """
        try:
            # Accedemos a la parte de numeración del documento
            numbering_part = doc.part.numbering_part
            if numbering_part is None:
                return False

            # 1. Buscar el elemento <w:num> que coincida con el numId del párrafo
            # Esto nos dará el abstractNumId
            num_element = numbering_part._element.find(
                f'.//w:num[@w:numId="{numId_val}"]',
                numbering_part._element.nsmap
            )
            if num_element is None:
                return False

            abstract_num_id_element = num_element.find(qn('w:abstractNumId'))
            if abstract_num_id_element is None:
                return False
            
            abstract_num_val = abstract_num_id_element.get(qn('w:val'))

            # 2. Buscar la definición abstracta <w:abstractNum> usando el ID encontrado
            abstract_num = numbering_part._element.find(
                f'.//w:abstractNum[@w:abstractNumId="{abstract_num_val}"]',
                numbering_part._element.nsmap
            )
            if abstract_num is None:
                return False

            # 3. Buscar el nivel 0 (<w:lvl w:ilvl="0">) para ver cómo está formateado
            lvl = abstract_num.find(
                f'.//w:lvl[@w:ilvl="0"]',
                numbering_part._element.nsmap
            )
            if lvl is None:
                return False

            # 4. Verificar el formato de número (<w:numFmt w:val="...">)
            num_fmt = lvl.find(qn('w:numFmt'))
            if num_fmt is None:
                return False

            fmt_val = num_fmt.get(qn('w:val'))
            
            # Si el formato es 'bullet', es lo que queremos eliminar.
            # Si es 'decimal', 'lowerLetter', etc., lo conservamos.
            return fmt_val == 'bullet'

        except Exception:
            # En caso de error de lectura XML, asumimos que NO es bullet para no romper nada importante
            return False

    def remove_bullets_keep_indent(self, file_path: str):
        try:
            doc = Document(file_path)
            modified = False
            
            for paragraph in doc.paragraphs:
                # 1. Obtener acceso al XML del párrafo
                p_pr = paragraph._element.find(qn('w:pPr'))
                
                if p_pr is not None:
                    # 2. Ver si tiene propiedades de numeración
                    num_pr = p_pr.find(qn('w:numPr'))
                    
                    if num_pr is not None:
                        # 3. Extraer el ID de la lista
                        num_id_element = num_pr.find(qn('w:numId'))
                        if num_id_element is not None:
                            val = num_id_element.get(qn('w:val'))
                            
                            # 4. INSPECCIÓN PROFUNDA:
                            # Solo procedemos si confirmamos que este ID es una VIÑETA.
                            if self._is_bullet_list(doc, val):
                                
                                # Capturar indentación actual antes de borrar
                                current_indent = paragraph.paragraph_format.left_indent
                                if current_indent is None:
                                    current_indent = Pt(36) # Default ~1.27cm
                                
                                # ELIMINAR LA VIÑETA (Borrar nodo numPr)
                                p_pr.remove(num_pr)
                                
                                # RESTAURAR INDENTACIÓN
                                paragraph.paragraph_format.left_indent = current_indent
                                
                                modified = True

            if modified:
                doc.save(file_path)
                print(f"    -> Post-procesamiento: Viñetas eliminadas (Numeración conservada) en {file_path}")
            else:
                print(f"    -> Post-procesamiento: No se requirieron cambios en {file_path}")

        except Exception as e:
            print(f"    -> ERROR en post-procesamiento: {e}")

    def _force_list_level_format(self, doc, num_id, level=1, format_type='upperLetter'):
        """
        Inspecciona el XML para encontrar la definición abstracta de la lista y 
        forzar el formato de número (ej. A, B, C) y EL INICIO EN 1 para un nivel específico.
        """
        try:
            numbering_part = doc.part.numbering_part
            if numbering_part is None:
                return

            # 1. Encontrar el elemento num (<w:num>) usando el numId actual
            num_element = numbering_part._element.find(
                f'.//w:num[@w:numId="{num_id}"]',
                numbering_part._element.nsmap
            )
            if num_element is None:
                return
            
            # Obtener el abstractNumId asociado
            abstract_num_id_element = num_element.find(qn('w:abstractNumId'))
            if abstract_num_id_element is None:
                return
            abstract_num_val = abstract_num_id_element.get(qn('w:val'))

            # 2. Encontrar la definición abstracta (<w:abstractNum>)
            abstract_num = numbering_part._element.find(
                f'.//w:abstractNum[@w:abstractNumId="{abstract_num_val}"]',
                numbering_part._element.nsmap
            )
            if abstract_num is None:
                return

            # 3. Modificar la Definición Abstracta (Afecta a todas las listas de este tipo)
            lvl = abstract_num.find(
                f'.//w:lvl[@w:ilvl="{level}"]',
                numbering_part._element.nsmap
            )
            if lvl is not None:
                # A. Forzar formato de letra (A., B., C.)
                num_fmt = lvl.find(qn('w:numFmt'))
                if num_fmt is not None:
                    num_fmt.set(qn('w:val'), format_type)
                
                # B. Forzar inicio en 1 (<w:start w:val="1"/>) - CORRECCIÓN CLAVE
                start_node = lvl.find(qn('w:start'))
                if start_node is not None:
                    start_node.set(qn('w:val'), '1')

            # 4. Modificar la Instancia Específica (Override) - CORRECCIÓN CLAVE
            # A veces la lista específica tiene una anulación (override) que fuerza el inicio en 4.
            # Debemos buscar <w:lvlOverride w:ilvl="1"> y cambiar <w:startOverride w:val="1">
            lvl_override = num_element.find(
                f'.//w:lvlOverride[@w:ilvl="{level}"]',
                numbering_part._element.nsmap
            )
            if lvl_override is not None:
                start_override = lvl_override.find(qn('w:startOverride'))
                if start_override is not None:
                    start_override.set(qn('w:val'), '1')

        except Exception as e:
            print(f"    -> Warning: No se pudo forzar el formato de lista: {e}")

    def convert_text_options_to_list(self, file_path: str):
        doc = Document(file_path)
        modified = False
        prefix_pattern = re.compile(r'^\s*([A-Da-d])(\.|\))\s+') 
        last_valid_num_id = None

        for paragraph in doc.paragraphs:
            # 1. Capturar contexto de lista
            p_pr = paragraph._element.find(qn('w:pPr'))
            if p_pr is not None:
                num_pr = p_pr.find(qn('w:numPr'))
                if num_pr is not None:
                    num_id_node = num_pr.find(qn('w:numId'))
                    if num_id_node is not None:
                        val_str = num_id_node.get(qn('w:val'))
                        ilvl_node = num_pr.find(qn('w:ilvl'))
                        if (ilvl_node is None or ilvl_node.get(qn('w:val')) == '0') and val_str:
                            last_valid_num_id = val_str

            # 2. Identificar si es una opción (A, B, C...)
            text = paragraph.text
            match = prefix_pattern.match(text)
            
            if match and last_valid_num_id:
                # Detectar si debe ser resaltada (basado en si algún run es negrita originalmente)
                is_correct_answer = any(run.bold for run in paragraph.runs)
                
                # --- LIMPIEZA TOTAL DE FORMATO HEREDADO ---
                if p_pr is not None:
                    p_style = p_pr.find(qn('w:pStyle'))
                    if p_style is not None:
                        p_pr.remove(p_style)

                # --- TRATAMIENTO QUIRÚRGICO DE RUNS ---
                chars_to_remove = len(match.group(0))
                for run in paragraph.runs:
                    # Quitamos la negrita de forma radical borrando la propiedad
                    run.bold = None 
                    
                    # Si es la respuesta correcta, aplicamos resaltado al texto
                    if is_correct_answer:
                        run.font.highlight_color = WD_COLOR_INDEX.YELLOW
                        # Acceso directo al XML para asegurar que no hay rPr de negrita
                        rPr = run._element.get_or_add_rPr()
                        for tag in [qn('w:b'), qn('w:bCs')]:
                            node = rPr.find(tag)
                            if node is not None:
                                rPr.remove(node)

                    # Eliminación del prefijo
                    if chars_to_remove > 0:
                        run_text = run.text
                        if len(run_text) <= chars_to_remove:
                            chars_to_remove -= len(run_text)
                            run.text = ""
                        else:
                            run.text = run_text[chars_to_remove:]
                            chars_to_remove = 0

                # --- RE-VINCULACIÓN A LA LISTA E INDENTACIÓN (PRIMERO) ---
                p_pr = paragraph._element.get_or_add_pPr()
                num_pr = p_pr.get_or_add_numPr()
                num_id_elem = num_pr.get_or_add_numId()
                num_id_elem.val = int(last_valid_num_id)
                ilvl_elem = num_pr.get_or_add_ilvl()
                ilvl_elem.val = 1 
                
                self._force_list_level_format(doc, last_valid_num_id, level=1, format_type='upperLetter')
                
                # Aplicamos la indentación ANTES de agregar el rPr manual
                paragraph.paragraph_format.left_indent = Pt(54)
                paragraph.paragraph_format.first_line_indent = Pt(-18)

                # --- RESALTADO DE LA MARCA DE PÁRRAFO / VIÑETA (AL FINAL) ---
                if is_correct_answer:
                    # Como ya se aplicó la indentación, el append colocará el rPr al final (Orden XML correcto)
                    r_pr = p_pr.find(qn('w:rPr'))
                    
                    if r_pr is None:
                        r_pr = OxmlElement('w:rPr')
                        p_pr.append(r_pr)
                    
                    highlight = r_pr.find(qn('w:highlight'))
                    if highlight is None:
                        highlight = OxmlElement('w:highlight')
                        r_pr.append(highlight)
                    
                    highlight.set(qn('w:val'), 'yellow')

                modified = True

        if modified:
            doc.save(file_path)

    def apply_margins(self, file_path: str, margin_cm: float = 1.0):
        """
        Ajusta los márgenes superior, inferior, izquierdo y derecho de todas
        las secciones del documento a la medida especificada en centímetros.
        """
        try:
            doc = Document(file_path)
            modified = False
            
            for section in doc.sections:
                section.top_margin = Cm(margin_cm)
                section.bottom_margin = Cm(margin_cm)
                section.left_margin = Cm(margin_cm)
                section.right_margin = Cm(margin_cm)
                modified = True
                
            if modified:
                doc.save(file_path)
                print(f"    -> Post-procesamiento: Márgenes ajustados a {margin_cm}cm en {file_path}")
                
        except Exception as e:
            print(f"    -> ERROR al ajustar márgenes en {file_path}: {e}")

    def restore_original_image_sizes(self, file_path: str, max_width_cm: float = 19.0):
        """
        Restaura el tamaño de las imágenes mediante inspección profunda del XML (OpenXML).
        Captura tanto <wp:inline> como <wp:anchor>, evadiendo los bloqueos de Pandoc y 
        encontrando las imágenes sin importar si están dentro de listas.
        """
        try:
            doc = Document(file_path)
            modified = False
            # 1 cm equivale a 360,000 EMUs (English Metric Units) en OpenXML
            max_width_emu = int(max_width_cm * 360000) 
            
            # Buscar TODOS los elementos de dibujo en el XML, estén donde estén
            for drawing in doc.element.xpath('.//w:drawing'):
                
                # 1. Identificar el contenedor principal (puede ser inline o anchor)
                container = drawing.find(qn('wp:inline'))
                if container is None:
                    container = drawing.find(qn('wp:anchor'))
                if container is None:
                    continue
                    
                # Extraer nodo de extensión principal
                wp_extent = container.find(qn('wp:extent'))
                if wp_extent is None:
                    continue
                    
                # 2. Buscar la referencia a la imagen física usando XPath (etiqueta a:blip)
                blips = drawing.xpath('.//a:blip')
                if not blips:
                    continue
                embed_id = blips[0].get(qn('r:embed'))
                if not embed_id:
                    continue
                    
                # 3. Extraer el archivo de imagen binario real del paquete DOCX
                try:
                    image_part = doc.part.related_parts[embed_id]
                    image_obj = image_part.image
                    
                    # Extraer píxeles y DPI nativos
                    px_width = image_obj.px_width
                    px_height = image_obj.px_height
                    
                    # Usar 96 DPI por defecto si la imagen no tiene meta-datos de densidad
                    dpi = getattr(image_obj, 'horz_dpi', 96)
                    if not dpi:
                        dpi = 96
                        
                    # Calcular el tamaño original en EMUs (1 pulgada = 914400 EMUs)
                    orig_width = int((px_width / dpi) * 914400)
                    orig_height = int((px_height / dpi) * 914400)
                    
                except Exception as e:
                    print(f"      -> Advertencia: No se pudo procesar la imagen {embed_id}: {e}")
                    continue
                    
                # 4. Calcular el tamaño físico preservando el margen de la página (ej. 19cm)
                if orig_width > max_width_emu:
                    ratio = max_width_emu / float(orig_width)
                    final_width = max_width_emu
                    final_height = int(orig_height * ratio)
                else:
                    final_width = orig_width
                    final_height = orig_height
                    
                # 5. Aplicar cambios forzados al XML
                current_width = int(wp_extent.get('cx', 0))
                if current_width != final_width:
                    # Sobrescribir contenedor principal
                    wp_extent.set('cx', str(final_width))
                    wp_extent.set('cy', str(final_height))
                    
                    # Es CRÍTICO sobrescribir también las etiquetas internas <a:ext> 
                    # para que Word no distorsione el aspecto de la imagen
                    a_exts = drawing.xpath('.//a:ext')
                    for ext in a_exts:
                        ext.set('cx', str(final_width))
                        ext.set('cy', str(final_height))
                        
                    modified = True

            if modified:
                doc.save(file_path)
                print(f"    -> Post-procesamiento: Imágenes XML restauradas con éxito en {file_path}")
            else:
                print(f"    -> Post-procesamiento: Las imágenes ya tenían su formato original en {file_path}")
                
        except Exception as e:
            print(f"    -> ERROR al restaurar imágenes por XML en {file_path}: {e}")